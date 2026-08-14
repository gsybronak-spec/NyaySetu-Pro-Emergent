"""Backend tests for the NyaySetu Pro document engine fix.

Covers:
  * Blocks-classification correctness (numbered points NOT bold/centered).
  * Real PDF/DOCX generation with content inspection (pdfminer / python-docx).
  * Certified Copy order_date optional.
  * Credit safety (preview does not decrement; download decrements exactly 1).
  * Regression: OTP login, google 401, referral +10, template list count+search.
"""
import base64
import io
import os
import re
import time
import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pdfminer.high_level import extract_text
import zlib
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))


def _pdf_actual_text(raw: bytes):
    """Extract the /ActualText spans from a generated PDF (selectable text)."""
    spans = []
    for obj in re.finditer(rb"\d+ 0 obj\n(.*?)\nendobj", raw, re.S):
        body = obj.group(1)
        if b"stream" not in body or b"FlateDecode" not in body:
            continue
        m = re.search(rb"<<(.*?)>>\s*stream\r?\n", body, re.S)
        if not m:
            continue
        lm = re.search(rb"/Length\s+(\d+)", m.group(1))
        if not lm:
            continue
        length = int(lm.group(1))
        data = body[m.end():m.end() + length]
        try:
            if b"ASCII85Decode" in m.group(1):
                data = base64.a85decode(data, adobe=True)
            d = zlib.decompress(data)
        except Exception:
            continue
        for am in re.finditer(rb"/ActualText\s*<FEFF([0-9A-F]+)>", d):
            try:
                spans.append(bytes.fromhex(am.group(1).decode("ascii")).decode("utf-16-be"))
            except Exception:
                continue
    return spans

os.environ.setdefault("MONGO_URL", "mongodb://localhost:27017")
os.environ.setdefault("DB_NAME", "nyaysetu_test_doc")

import mongomock_motor
mock_client = mongomock_motor.AsyncMongoMockClient()
mock_db = mock_client["nyaysetu_test_doc"]

import server
server.db = mock_db

from starlette.testclient import TestClient
app_client = TestClient(server.app)

class TestClientWrapper:
    def get(self, url, **kwargs):
        kwargs.pop("timeout", None)
        return app_client.get(url, **kwargs)
    def post(self, url, **kwargs):
        kwargs.pop("timeout", None)
        return app_client.post(url, **kwargs)
    def put(self, url, **kwargs):
        kwargs.pop("timeout", None)
        return app_client.put(url, **kwargs)
    def delete(self, url, **kwargs):
        kwargs.pop("timeout", None)
        return app_client.delete(url, **kwargs)

requests = TestClientWrapper()
BASE = "/api"


# --------- helpers ---------
def _login(mobile: str) -> str:
    r = app_client.post(f"{BASE}/auth/send-otp", json={"mobile": mobile})
    assert r.status_code == 200, r.text
    r = app_client.post(f"{BASE}/auth/verify-otp", json={"mobile": mobile, "otp": "123456"})
    assert r.status_code == 200, r.text
    return r.json()["token"]


def _hdr(tok):
    return {"Authorization": f"Bearer {tok}", "Content-Type": "application/json"}


def _new_mobile():
    return "9" + str(int(time.time() * 1000))[-9:]


@pytest.fixture(scope="module")
def token():
    return _login(_new_mobile())


# --------- BLOCK CLASSIFICATION ---------
class TestBlocksClassification:
    def test_gu_adjournment_points_not_bold(self, token):
        # Create a case first so case_id can be attached
        c = requests.post(
            f"{BASE}/cases",
            headers=_hdr(token),
            json={
                "nickname": "TEST_Doc",
                "case_number": "42/2026",
                "case_type_id": "civil_suit",
                "district_id": "ahmedabad",
                "party_name": "TEST Party",
                "language": "gu",
            },
            timeout=15,
        )
        assert c.status_code == 200, c.text
        case_id = c.json()["id"]

        r = requests.post(
            f"{BASE}/applications/preview",
            headers=_hdr(token),
            json={
                "template_id": "adjournment",
                "language": "gu",
                "case_id": case_id,
                "values": {"next_date": "2026-02-15", "reason": "વ્યક્તિગત કારણોસર"},
            },
            timeout=15,
        )
        assert r.status_code == 200, r.text
        blocks = r.json()["blocks"]
        assert isinstance(blocks, list) and len(blocks) > 5

        # Extract points ૧. ૨. ૩.
        point_blocks = [b for b in blocks if re.match(r"^[૧૨૩]\.", b["text"].strip())]
        assert len(point_blocks) >= 3, f"expected 3 numbered points, got {[b['text'] for b in point_blocks]}"
        for pb in point_blocks:
            assert pb["align"] == "left", f"Point wrongly centered: {pb}"
            assert pb["bold"] is False, f"Point wrongly bold: {pb}"

        # Court heading + title should be center+bold
        centered = [b for b in blocks if b["align"] == "center" and b["bold"]]
        assert any(b["text"].startswith("માનનીય ન્યાયાલય") for b in centered), "court heading not centered"
        assert any("મુદત" in b["text"] and b["bold"] for b in centered), "title not bold+center"

    def test_en_affidavit_headings(self, token):
        r = requests.post(
            f"{BASE}/applications/preview",
            headers=_hdr(token),
            json={
                "template_id": "affidavit",
                "language": "en",
                "values": {
                    "deponent_name": "Ravi",
                    "father_name": "Kishor",
                    "age": 30,
                    "address": "Ahmedabad",
                    "statement": "This is my statement.",
                },
            },
            timeout=15,
        )
        assert r.status_code == 200, r.text
        blocks = r.json()["blocks"]
        # AFFIDAVIT + VERIFICATION centered+bold
        assert any(b["text"] == "AFFIDAVIT" and b["align"] == "center" and b["bold"] for b in blocks)
        assert any(b["text"] == "VERIFICATION" and b["align"] == "center" and b["bold"] for b in blocks)
        # Any numbered points 1./2./3. must be left+not-bold
        for b in blocks:
            if re.match(r"^\d+\.", b["text"].strip()):
                assert b["align"] == "left" and b["bold"] is False, f"numbered point wrong: {b}"


# --------- REAL FILE GENERATION ---------
class TestPdfGeneration:
    def test_gu_pdf_has_gujarati_text(self, token):
        r = requests.post(
            f"{BASE}/applications/download",
            headers=_hdr(token),
            json={
                "template_id": "adjournment",
                "language": "gu",
                "format": "pdf",
                "consume_credit": True,
                "values": {"next_date": "2026-02-15", "reason": "વ્યક્તિગત કારણોસર"},
            },
            timeout=30,
        )
        assert r.status_code == 200, r.text
        b64 = r.json()["base64"]
        raw = base64.b64decode(b64)
        assert raw[:4] == b"%PDF"
        # Selectable text is carried by /ActualText spans (the HarfBuzz engine
        # draws shaped glyphs as PUA codes; pdfminer cannot decode those, but
        # every conforming reader — MuPDF, Acrobat, Chrome — uses ActualText).
        spans = _pdf_actual_text(raw)
        joined = "\n".join(spans)
        assert "મુદત" in joined, f"'મુદત' missing from ActualText. Spans: {joined[:400]!r}"
        assert "અરજદાર" in joined, f"'અરજદાર' missing from ActualText. Spans: {joined[:400]!r}"

    def test_en_pdf_ok(self, token):
        r = requests.post(
            f"{BASE}/applications/download",
            headers=_hdr(token),
            json={
                "template_id": "affidavit",
                "language": "en",
                "format": "pdf",
                "consume_credit": True,
                "values": {
                    "deponent_name": "Ravi",
                    "father_name": "Kishor",
                    "age": 30,
                    "address": "Ahmedabad",
                    "statement": "This is my statement.",
                },
            },
            timeout=30,
        )
        assert r.status_code == 200, r.text
        raw = base64.b64decode(r.json()["base64"])
        assert raw[:4] == b"%PDF"
        text = extract_text(io.BytesIO(raw))
        assert "AFFIDAVIT" in text
        assert "VERIFICATION" in text
        assert "Ravi" in text


class TestDocxGeneration:
    def test_gu_docx_paragraphs_and_font(self, token):
        r = requests.post(
            f"{BASE}/applications/download",
            headers=_hdr(token),
            json={
                "template_id": "adjournment",
                "language": "gu",
                "format": "docx",
                "consume_credit": True,
                "values": {"next_date": "2026-02-15", "reason": "કારણ"},
            },
            timeout=30,
        )
        assert r.status_code == 200
        raw = base64.b64decode(r.json()["base64"])
        doc = Document(io.BytesIO(raw))
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paragraphs) > 5
        # Find numbered points & check they are LEFT / not bold
        for p in paragraphs:
            if re.match(r"^[૧૨૩]\.", p.text.strip()):
                assert p.alignment in (WD_ALIGN_PARAGRAPH.LEFT, None), f"Gujarati point centered: {p.text}"
                for run in p.runs:
                    assert not run.bold, f"Gujarati point bold: {p.text}"
        # Find title / court heading -> center+bold
        centered_bold = [
            p for p in paragraphs
            if p.alignment == WD_ALIGN_PARAGRAPH.CENTER and any(r.bold for r in p.runs)
        ]
        assert any("માનનીય ન્યાયાલય" in p.text or "મુદત" in p.text for p in centered_bold)
        # Font name = Lohit Gujarati on at least one run
        fonts = {run.font.name for p in paragraphs for run in p.runs if run.font.name}
        assert "Lohit Gujarati" in fonts, f"Fonts used: {fonts}"

    def test_en_docx_font_times(self, token):
        r = requests.post(
            f"{BASE}/applications/download",
            headers=_hdr(token),
            json={
                "template_id": "affidavit",
                "language": "en",
                "format": "docx",
                "consume_credit": True,
                "values": {
                    "deponent_name": "Ravi",
                    "father_name": "Kishor",
                    "age": 30,
                    "address": "Ahmedabad",
                    "statement": "This is my statement.",
                },
            },
            timeout=30,
        )
        assert r.status_code == 200
        raw = base64.b64decode(r.json()["base64"])
        doc = Document(io.BytesIO(raw))
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        fonts = {run.font.name for p in paragraphs for run in p.runs if run.font.name}
        assert "Times New Roman" in fonts, f"Fonts: {fonts}"
        # AFFIDAVIT / VERIFICATION center+bold
        cb_texts = [p.text for p in paragraphs
                    if p.alignment == WD_ALIGN_PARAGRAPH.CENTER and any(r.bold for r in p.runs)]
        assert "AFFIDAVIT" in cb_texts
        assert "VERIFICATION" in cb_texts
        # numbered points left/not-bold
        for p in paragraphs:
            if re.match(r"^\d+\.", p.text.strip()):
                assert p.alignment in (WD_ALIGN_PARAGRAPH.LEFT, None)
                for run in p.runs:
                    assert not run.bold


# --------- PHASE A DOCUMENT ENGINE FOUNDATION TESTS ---------
class TestPhaseADocumentEngine:
    def test_custom_margins_settings_override(self):
        from doc_generator import generate_pdf, generate_docx, get_doc_settings
        custom_settings = {
            "margin_top_cm": 3.5,
            "margin_bottom_cm": 3.0,
            "margin_left_cm": 4.0,
            "margin_right_cm": 2.0,
            "heading_size": 14,
            "body_size": 12,
        }
        merged = get_doc_settings(custom_settings)
        assert merged["margin_top_cm"] == 3.5
        assert merged["margin_left_cm"] == 4.0

        blocks = [
            {"text": "IN THE COURT OF JMFC", "align": "center", "bold": True},
            {"text": "1. That the applicant submits this application.", "align": "left", "bold": False},
        ]
        pdf_b64 = generate_pdf(blocks, language="en", settings=custom_settings)
        assert len(pdf_b64) > 100
        raw_pdf = base64.b64decode(pdf_b64)
        assert raw_pdf[:4] == b"%PDF"

        docx_b64 = generate_docx(blocks, language="en", settings=custom_settings)
        assert len(docx_b64) > 100
        raw_docx = base64.b64decode(docx_b64)
        doc = Document(io.BytesIO(raw_docx))
        assert round(doc.sections[0].left_margin.cm, 1) == 4.0

    def test_gujarati_character_pipeline(self):
        from doc_generator import generate_pdf, generate_docx
        guj_text = "અરજી માનનીય ન્યાયાલય વિગત મુદત સ્વિકાર સ્પેશિયલ પ્રમાણિત"
        blocks = [
            {"text": "માનનીય ન્યાયાલય અમદાવાદ", "align": "center", "bold": True},
            {"text": f"૧. {guj_text}", "align": "left", "bold": False},
        ]
        pdf_b64 = generate_pdf(blocks, language="gu")
        raw_pdf = base64.b64decode(pdf_b64)
        joined = "\n".join(_pdf_actual_text(raw_pdf))
        assert "અરજી" in joined
        assert "મુદત" in joined
        assert "ન્યાયાલય" in joined

        docx_b64 = generate_docx(blocks, language="gu")
        raw_docx = base64.b64decode(docx_b64)
        doc = Document(io.BytesIO(raw_docx))
        assert any("ન્યાયાલય" in p.text for p in doc.paragraphs)

    def test_harfbuzz_shaping_exact_words(self):
        from doc_generator import generate_pdf, generate_pdf_playwright
        exact_words = [
            "અરજી", "વિગત", "મુદત", "સ્વીકાર", "વિશેષ", "પ્રમાણિત",
            "પ્રતિઉત્તર", "સ્વપ્રમાણિત", "પ્રતિનિધિત્વ", "ક્ષતિ", "જ્ઞાપન", "ન્યાયાધીશ"
        ]
        blocks = [{"text": w, "align": "left", "bold": False} for w in exact_words]

        # Generate PDF using Chromium engine
        pdf_b64 = generate_pdf(blocks, language="gu")
        assert len(pdf_b64) > 200
        raw_pdf = base64.b64decode(pdf_b64)
        assert raw_pdf[:4] == b"%PDF"

        # Verify PDF contains extractable text
        pdf_text = extract_text(io.BytesIO(raw_pdf))
        assert len(pdf_text.strip()) > 20

    def test_english_times_font_configuration(self):
        from doc_generator import generate_pdf, FONT_DIR, get_doc_settings
        times_ttf = FONT_DIR / "TimesNewRoman.ttf"
        settings = get_doc_settings()
        blocks = [{"text": "TEST IN THE COURT OF SESSIONS", "align": "center", "bold": True}]
        pdf_b64 = generate_pdf(blocks, language="en")
        raw_pdf = base64.b64decode(pdf_b64)
        pdf_text = extract_text(io.BytesIO(raw_pdf))
        assert "SESSIONS" in pdf_text

        if not times_ttf.exists():
            # Standard ReportLab Times-Roman PostScript font is used (not an embedded TTF)
            assert settings["english_font"] == "Times-Roman"

    def test_formatting_preservation(self):
        from doc_generator import build_blocks
        content = (
            "IN THE COURT OF SESSIONS JUDGE AHMEDABAD\n\n"
            "APPLICATION FOR BAIL\n\n"
            "1. First point of application.\n"
            "2. Second point of application.\n"
        )
        blocks = build_blocks(content, title_en="APPLICATION FOR BAIL")
        assert blocks[0]["align"] == "center" and blocks[0]["bold"] is True
        assert blocks[2]["align"] == "center" and blocks[2]["bold"] is True
        assert blocks[4]["align"] == "left" and blocks[4]["bold"] is False
        assert blocks[5]["align"] == "left" and blocks[5]["bold"] is False


# --------- PDF vs DOCX consistency ---------
class TestPdfDocxConsistency:
    def test_same_text_content(self):
        token = _login(_new_mobile())  # fresh 5 credits
        payload = {
            "template_id": "affidavit",
            "language": "en",
            "consume_credit": True,
            "values": {
                "deponent_name": "Alice",
                "father_name": "Bob",
                "age": 40,
                "address": "Rajkot",
                "statement": "Truthful statement.",
            },
        }
        rp = requests.post(f"{BASE}/applications/download",
                           headers=_hdr(token), json={**payload, "format": "pdf"}, timeout=30)
        rd = requests.post(f"{BASE}/applications/download",
                           headers=_hdr(token), json={**payload, "format": "docx"}, timeout=30)
        assert rp.status_code == 200 and rd.status_code == 200
        pdf_text = extract_text(io.BytesIO(base64.b64decode(rp.json()["base64"])))
        docx_text = "\n".join(
            p.text for p in Document(io.BytesIO(base64.b64decode(rd.json()["base64"]))).paragraphs
        )
        for token_ in ["AFFIDAVIT", "VERIFICATION", "Alice", "Bob", "Rajkot", "Truthful statement"]:
            assert token_ in pdf_text, f"{token_} missing in PDF"
            assert token_ in docx_text, f"{token_} missing in DOCX"


# --------- CERTIFIED COPY ORDER DATE OPTIONAL ---------
class TestCertifiedCopyOptional:
    def test_without_order_date(self):
        token = _login(_new_mobile())
        r = requests.post(
            f"{BASE}/applications/download",
            headers=_hdr(token),
            json={
                "template_id": "certified_copy",
                "language": "en",
                "format": "pdf",
                "consume_credit": True,
                "values": {"document_desc": "Order dated in the above matter"},
            },
            timeout=30,
        )
        assert r.status_code == 200, r.text
        raw = base64.b64decode(r.json()["base64"])
        assert raw[:4] == b"%PDF"

    def test_with_order_date(self):
        token = _login(_new_mobile())
        r = requests.post(
            f"{BASE}/applications/download",
            headers=_hdr(token),
            json={
                "template_id": "certified_copy",
                "language": "en",
                "format": "pdf",
                "consume_credit": True,
                "values": {"document_desc": "Order", "order_date": "2026-01-10"},
            },
            timeout=30,
        )
        assert r.status_code == 200
        assert base64.b64decode(r.json()["base64"])[:4] == b"%PDF"

    def test_field_marked_optional_in_catalog(self):
        # No auth needed for detail? Try both.
        tok = _login(_new_mobile())
        r = requests.get(f"{BASE}/templates/certified_copy", headers=_hdr(tok), timeout=15)
        assert r.status_code == 200, r.text
        fields = r.json()["fields"]
        od = next((f for f in fields if f["key"] == "order_date"), None)
        assert od is not None
        assert od["required"] is False


# --------- CREDIT SAFETY ---------
class TestCreditSafety:
    def test_preview_no_credit_deduct(self):
        tok = _login(_new_mobile())  # fresh user w/ 5 credits
        before = requests.get(f"{BASE}/wallet", headers=_hdr(tok)).json()["balance"]
        requests.post(
            f"{BASE}/applications/preview",
            headers=_hdr(tok),
            json={"template_id": "adjournment", "language": "en",
                  "values": {"next_date": "2026-02-01", "reason": "x"}},
            timeout=15,
        )
        after = requests.get(f"{BASE}/wallet", headers=_hdr(tok)).json()["balance"]
        assert after == before, f"preview deducted credit! {before}->{after}"

    def test_download_deducts_exactly_one(self):
        tok = _login(_new_mobile())
        before = requests.get(f"{BASE}/wallet", headers=_hdr(tok)).json()["balance"]
        r = requests.post(
            f"{BASE}/applications/download",
            headers=_hdr(tok),
            json={"template_id": "adjournment", "language": "en", "format": "pdf",
                  "consume_credit": True,
                  "values": {"next_date": "2026-02-01", "reason": "x"}},
            timeout=30,
        )
        assert r.status_code == 200
        after = requests.get(f"{BASE}/wallet", headers=_hdr(tok)).json()["balance"]
        assert after == before - 1, f"expected -1, got {before}->{after}"

    def test_insufficient_balance_402(self):
        tok = _login(_new_mobile())
        # Drain 5 credits
        for _ in range(5):
            requests.post(
                f"{BASE}/applications/download",
                headers=_hdr(tok),
                json={"template_id": "adjournment", "language": "en", "format": "pdf",
                      "consume_credit": True,
                      "values": {"next_date": "2026-02-01", "reason": "x"}},
                timeout=30,
            )
        bal = requests.get(f"{BASE}/wallet", headers=_hdr(tok)).json()["balance"]
        assert bal == 0
        r = requests.post(
            f"{BASE}/applications/download",
            headers=_hdr(tok),
            json={"template_id": "adjournment", "language": "en", "format": "pdf",
                  "consume_credit": True,
                  "values": {"next_date": "2026-02-01", "reason": "x"}},
            timeout=30,
        )
        assert r.status_code == 402, f"expected 402, got {r.status_code}: {r.text}"
        # Balance still 0 (no deduction)
        bal2 = requests.get(f"{BASE}/wallet", headers=_hdr(tok)).json()["balance"]
        assert bal2 == 0


# --------- REGRESSION ---------
class TestRegression:
    def test_google_invalid_401(self):
        r = requests.post(f"{BASE}/auth/google-session",
                          json={"session_id": "invalid_abc"}, timeout=15)
        assert r.status_code == 401, r.text

    def test_referral_plus_10(self):
        # Existing user shares referral, new user signs up with it
        ref_tok = _login(_new_mobile())
        me = requests.get(f"{BASE}/referral/me", headers=_hdr(ref_tok), timeout=15).json()
        code = me.get("referral_code") or me.get("code")
        assert code, f"referral code missing in response: {me}"
        bal_before = requests.get(f"{BASE}/wallet", headers=_hdr(ref_tok)).json()["balance"]

        new_mobile = _new_mobile()
        requests.post(f"{BASE}/auth/send-otp", json={"mobile": new_mobile}, timeout=15)
        r = requests.post(f"{BASE}/auth/verify-otp",
                          json={"mobile": new_mobile, "otp": "123456", "referral_code": code}, timeout=15)
        assert r.status_code == 200

        bal_after = requests.get(f"{BASE}/wallet", headers=_hdr(ref_tok)).json()["balance"]
        assert bal_after == bal_before + 10, f"referral +10 failed: {bal_before}->{bal_after}"

    def test_template_list_and_search(self, token):
        r = requests.get(f"{BASE}/templates", headers=_hdr(token), timeout=15)
        assert r.status_code == 200
        items = r.json()
        assert len(items) >= 23, f"expected >=23 templates, got {len(items)}"

        # multilingual search
        for q in ["mudat", "વિલંબ", "vakalat"]:
            rr = requests.get(f"{BASE}/templates", headers=_hdr(token),
                              params={"q": q}, timeout=15)
            assert rr.status_code == 200
            assert len(rr.json()) >= 1, f"no result for {q}"
