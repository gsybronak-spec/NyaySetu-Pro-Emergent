"""Tests for the Admin Word Template Import feature + publish-guard fix.

Covers:
- Analyze: valid .docx -> page size / margins / fonts / fields / placeholders
- Analyze: A4 and Legal page sizes detected from section dimensions
- Analyze: non-docx file names and corrupt content rejected (400, readable)
- Analyze: unauthenticated 401, lawyer 401, regular admin 403, super admin OK
- Create: draft template created (source=imported), not visible to lawyers
- Create: unknown placeholders -> readable 400 validation error
- Create: duplicate id -> 409
- Publish: imported draft becomes visible in the lawyer template API
- Exactly one template record exists for an imported id (no duplicates)
- Existing templates remain unchanged
- Publish guard: malformed partial record -> readable 400, never a 500
- Generated PDF uses the configured page size (A4)

Uses mongomock_motor (same pattern as the existing suite).
"""

import base64
import io
import os
import re
import sys
import uuid
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
os.environ.setdefault("MONGO_URL", "mongodb://localhost:27017")
os.environ.setdefault("DB_NAME", "nyaysetu_test_word_import")

import pytest
import pytest_asyncio
import bcrypt
from datetime import datetime, timezone

import mongomock_motor
mock_client = mongomock_motor.AsyncMongoMockClient()
mock_db = mock_client["nyaysetu_test_word_import"]

import server
server.db = mock_db
db = mock_db
app = server.app

from server import make_admin_token, make_token, now
from seed_data import TEMPLATES
from httpx import AsyncClient, ASGITransport

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt

API = "/api"
ADMIN = "/api/admin"

SEED_IDS = {t["id"] for t in TEMPLATES}
SEED_COUNT = len(SEED_IDS)

COLLECTIONS = ["admin_users", "users", "wallets", "cases", "drafts",
               "applications", "transactions", "referrals",
               "templates", "template_versions", "otps", "audit_logs",
               "plans", "settings"]


def make_sample_docx(page_size: str = "A4", margins=(2.0, 2.0, 4.0, 4.0)) -> bytes:
    """Build a two-part Word document: page 1 field spec, page 2 legal draft."""
    doc = Document()
    sec = doc.sections[0]
    if page_size == "Legal":
        sec.page_width = Cm(21.59)
        sec.page_height = Cm(35.56)
    else:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
    sec.top_margin = Cm(margins[0])
    sec.bottom_margin = Cm(margins[1])
    sec.left_margin = Cm(margins[2])
    sec.right_margin = Cm(margins[3])

    # ---- Page 1: field definitions (spec) ----
    doc.add_paragraph("અરજદારનું નામ [________]")
    doc.add_paragraph("કોર્ટ નુ નામ (ડ્રોપ બોક્ષ)")
    doc.add_paragraph("0 ફરીયાદી 0 અરજદાર 0 વાદી")
    doc.add_paragraph("તારીખ : [__ / __ / 20__]")

    # ---- explicit page break -> page 2 ----
    br = doc.add_paragraph()
    br.add_run().add_break(WD_BREAK.PAGE)

    # ---- Page 2: the actual legal draft ----
    title = doc.add_paragraph("દસ્તાવેજ પરત મેળવવાની અરજી")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    body = doc.add_paragraph("અરજદારનું નામ : {{applicant_name}}")
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("કોર્ટ : {{court}}")
    doc.add_paragraph("કેસ નંબર : {{case_number}}")
    doc.add_paragraph("અરજદાર : {{applicant_role}}")
    doc.add_paragraph("સામાવાળા : {{opposite_role}}")
    doc.add_paragraph("કેસ સ્ટેટસ : {{case_status}}")
    sig = doc.add_paragraph("અરજદારના એડવોકેટ")
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def b64(data: bytes) -> str:
    return base64.b64encode(data).decode("ascii")


@pytest_asyncio.fixture(scope="function")
async def client():
    server.db = mock_db
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        yield c


@pytest_asyncio.fixture(scope="function")
async def clean_db():
    for coll in COLLECTIONS:
        await db[coll].drop()
    yield
    for coll in COLLECTIONS:
        await db[coll].drop()


async def create_admin(role="super_admin"):
    admin_id = str(uuid.uuid4())
    hashed = bcrypt.hashpw(b"TestPass123!", bcrypt.gensalt()).decode("utf-8")
    admin = {
        "id": admin_id,
        "email": f"{role}@test.com",
        "password_hash": hashed,
        "name": "Test Admin",
        "role": role,
        "active": True,
        "last_login": None,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.admin_users.insert_one(admin.copy())
    token = make_admin_token(admin_id, admin["email"], admin["role"])
    return admin, token


async def create_lawyer():
    user_id = str(uuid.uuid4())
    mobile = "98" + str(uuid.uuid4().int)[:8]
    user = {
        "id": user_id,
        "mobile": mobile,
        "name": "Test Lawyer",
        "active": True,
        "created_at": now().isoformat(),
        "updated_at": now().isoformat(),
    }
    await db.users.insert_one(user.copy())
    await db.wallets.insert_one({"user_id": user_id, "balance": 10, "total_used": 0,
                                "created_at": now().isoformat(), "updated_at": now().isoformat()})
    token = make_token(user_id)
    return user, token


def H(token):
    return {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}


def field(key, label_en, label_gu, ftype="text", required=True, options=None):
    return {
        "key": key,
        "label_en": label_en,
        "label_gu": label_gu,
        "type": ftype,
        "required": required,
        "order": 0,
        "default_value": None,
        "options": options or [],
        "validation": None,
    }


def reviewed_fields():
    """Simulate the admin review of the extracted fields (keys linked to placeholders)."""
    return [
        field("applicant_name", "Applicant Name", "અરજદારનું નામ"),
        field("court", "Court", "કોર્ટ", ftype="select"),
        field("case_number", "Case Number", "કેસ નંબર"),
        field("applicant_role", "Applicant Role", "અરજદાર", ftype="radio",
              options=[{"value": "ફરીયાદી", "label_en": "ફરીયાદી", "label_gu": "ફરીયાદી"},
                       {"value": "અરજદાર", "label_en": "અરજદાર", "label_gu": "અરજદાર"}]),
        field("opposite_role", "Opposite Role", "સામાવાળા", ftype="radio",
              options=[{"value": "વાદી", "label_en": "વાદી", "label_gu": "વાદી"},
                       {"value": "પ્રતિવાદી", "label_en": "પ્રતિવાદી", "label_gu": "પ્રતિવાદી"}]),
        field("case_status", "Case Status", "કેસ સ્ટેટસ", ftype="radio",
              options=[{"value": "ચાલુ", "label_en": "ચાલુ", "label_gu": "ચાલુ"},
                       {"value": "ડિસ્પોસ્ડ", "label_en": "ડિસ્પોસ્ડ", "label_gu": "ડિસ્પોસ્ડ"}]),
    ]


CONTENT_GU = (
    "દસ્તાવેજ પરત મેળવવાની અરજી\n"
    "અરજદારનું નામ : {{applicant_name}}\n"
    "કોર્ટ : {{court}}\n"
    "કેસ નંબર : {{case_number}}\n"
    "અરજદાર : {{applicant_role}}\n"
    "સામાવાળા : {{opposite_role}}\n"
    "કેસ સ્ટેટસ : {{case_status}}\n"
    "અરજદારના એડવોકેટ"
)


def settings_from_analysis(a):
    s = dict(a["settings"])
    return s


# ---------------------------------------------------------------------------
# Analyze
# ---------------------------------------------------------------------------

async def test_analyze_valid_docx(client, clean_db):
    admin, token = await create_admin()
    data = make_sample_docx()
    r = await client.post(f"{ADMIN}/templates/import-word/analyze",
                          json={"file_name": "sample.docx", "content_base64": b64(data)},
                          headers=H(token))
    assert r.status_code == 200, r.text
    a = r.json()
    assert a["page_size"] == "A4"
    assert a["margins_cm"]["top_cm"] == 2.0
    assert a["margins_cm"]["bottom_cm"] == 2.0
    assert a["margins_cm"]["left_cm"] == 4.0
    assert a["margins_cm"]["right_cm"] == 4.0
    assert a["page_break_detected"] is True
    # Placeholders from the draft
    for ph in ("applicant_name", "court", "case_number", "applicant_role", "opposite_role", "case_status"):
        assert ph in a["placeholders"], a["placeholders"]
    # Draft content preserved (Gujarati + placeholders)
    assert "દસ્તાવેજ પરત મેળવવાની અરજી" in a["draft_content_gu"]
    assert "{{applicant_name}}" in a["draft_content_gu"]
    # Spec page detection
    keys = {f["key"] for f in a["fields"]}
    types = {f["type"] for f in a["fields"]}
    assert "radio" in types  # radio glyph line detected
    assert "select" in types  # dropdown marker detected
    assert "date" in types  # date marker detected
    # Fonts
    assert a["fonts"]["body_size"] >= 10
    assert a["settings"]["page_size"] == "A4"
    assert a["settings"]["margin_left_cm"] == 4.0


async def test_analyze_legal_page_size(client, clean_db):
    admin, token = await create_admin()
    r = await client.post(f"{ADMIN}/templates/import-word/analyze",
                          json={"file_name": "affidavit.docx", "content_base64": b64(make_sample_docx(page_size="Legal"))},
                          headers=H(token))
    assert r.status_code == 200, r.text
    assert r.json()["page_size"] == "Legal"


async def test_analyze_rejects_non_docx(client, clean_db):
    admin, token = await create_admin()
    r = await client.post(f"{ADMIN}/templates/import-word/analyze",
                          json={"file_name": "notes.pdf", "content_base64": b64(b"%PDF-1.4 fake")},
                          headers=H(token))
    assert r.status_code == 400, r.text
    assert "Only .docx" in r.json()["detail"]


async def test_analyze_rejects_corrupt_docx(client, clean_db):
    admin, token = await create_admin()
    r = await client.post(f"{ADMIN}/templates/import-word/analyze",
                          json={"file_name": "broken.docx", "content_base64": b64(b"not a zip file at all")},
                          headers=H(token))
    assert r.status_code == 400, r.text
    assert "could not be parsed" in r.json()["detail"]


async def test_analyze_auth_zones(client, clean_db):
    admin, token = await create_admin()
    # Unauthenticated -> 401
    r = await client.post(f"{ADMIN}/templates/import-word/analyze",
                          json={"file_name": "x.docx", "content_base64": b64(make_sample_docx())})
    assert r.status_code == 401, r.text
    # Lawyer -> 401
    _, ltoken = await create_lawyer()
    r = await client.post(f"{ADMIN}/templates/import-word/analyze",
                          json={"file_name": "x.docx", "content_base64": b64(make_sample_docx())},
                          headers=H(ltoken))
    assert r.status_code == 401, r.text
    # Regular admin -> 403
    _, atoken = await create_admin(role="admin")
    r = await client.post(f"{ADMIN}/templates/import-word/analyze",
                          json={"file_name": "x.docx", "content_base64": b64(make_sample_docx())},
                          headers=H(atoken))
    assert r.status_code == 403, r.text
    # Super admin -> 200
    r = await client.post(f"{ADMIN}/templates/import-word/analyze",
                          json={"file_name": "x.docx", "content_base64": b64(make_sample_docx())},
                          headers=H(token))
    assert r.status_code == 200, r.text
    # Audit trail exists for the successful analyze
    logs = await db.audit_logs.find({"action": "template_import_analyze"}).to_list(10)
    assert len(logs) >= 1


# ---------------------------------------------------------------------------
# Create (draft) + publish + lawyer visibility
# ---------------------------------------------------------------------------

async def test_import_create_and_publish_flow(client, clean_db):
    admin, token = await create_admin()
    a = (await client.post(f"{ADMIN}/templates/import-word/analyze",
                           json={"file_name": "sample.docx", "content_base64": b64(make_sample_docx())},
                           headers=H(token))).json()

    create_payload = {
        "id": "test_return_documents_import",
        "name_en": "Application for Return of Documents (Import)",
        "name_gu": "દસ્તાવેજ પરત મેળવવાની અરજી",
        "category": "Civil",
        "description": "Imported from Word document: sample.docx",
        "content_gu": CONTENT_GU,
        "fields": reviewed_fields(),
        "settings": settings_from_analysis(a),
    }
    r = await client.post(f"{ADMIN}/templates/import-word", json=create_payload, headers=H(token))
    assert r.status_code == 200, r.text
    created = r.json()
    assert created["status"] == "draft"
    assert created["source"] == "imported"
    assert created["id"] == "test_return_documents_import"

    # Exactly one DB record for this id
    n = await db.templates.count_documents({"id": "test_return_documents_import"})
    assert n == 1

    # Not visible to lawyers while draft
    r = await client.get(f"{API}/templates/test_return_documents_import")
    assert r.status_code == 404

    # Draft appears in the admin list
    r = await client.get(f"{ADMIN}/templates", headers=H(token))
    assert r.status_code == 200
    ids = {t["id"] for t in r.json()}
    assert "test_return_documents_import" in ids

    # Publish (super admin) -> visible to lawyers
    r = await client.post(f"{ADMIN}/templates/test_return_documents_import/publish", headers=H(token))
    assert r.status_code == 200, r.text
    r = await client.get(f"{API}/templates/test_return_documents_import")
    assert r.status_code == 200
    t = r.json()
    assert t["fields"][0]["key"] == "applicant_name"
    # Radio options preserved
    cs = next(f for f in t["fields"] if f["key"] == "case_status")
    assert [o["value"] for o in cs["options"]] == ["ચાલુ", "ડિસ્પોસ્ડ"]
    # Gujarati content intact
    assert "દસ્તાવેજ પરત મેળવવાની અરજી" in t["content_gu"]

    # Still exactly one record
    n = await db.templates.count_documents({"id": "test_return_documents_import"})
    assert n == 1
    # Seed templates untouched (no duplicates of any seed id)
    for sid in SEED_IDS:
        assert await db.templates.count_documents({"id": sid}) <= 1
    # Audit trail
    logs = await db.audit_logs.find({"action": "template_import"}).to_list(10)
    assert len(logs) == 1
    assert logs[0]["target"] == "test_return_documents_import"


async def test_import_create_unknown_placeholder(client, clean_db):
    admin, token = await create_admin()
    a = (await client.post(f"{ADMIN}/templates/import-word/analyze",
                           json={"file_name": "sample.docx", "content_base64": b64(make_sample_docx())},
                           headers=H(token))).json()
    payload = {
        "id": "bad_placeholder_import",
        "name_en": "Bad Import",
        "name_gu": "બેડ",
        "category": "General",
        "content_gu": "નામ : {{not_a_field}}",
        "fields": [],
        "settings": settings_from_analysis(a),
    }
    r = await client.post(f"{ADMIN}/templates/import-word", json=payload, headers=H(token))
    assert r.status_code == 400, r.text
    assert "not_a_field" in r.json()["detail"]


async def test_import_create_duplicate_id(client, clean_db):
    admin, token = await create_admin()
    a = (await client.post(f"{ADMIN}/templates/import-word/analyze",
                           json={"file_name": "sample.docx", "content_base64": b64(make_sample_docx())},
                           headers=H(token))).json()
    payload = {
        "id": "dup_import",
        "name_en": "Dup",
        "name_gu": "ડુપ",
        "content_gu": CONTENT_GU,
        "fields": reviewed_fields(),
        "settings": settings_from_analysis(a),
    }
    r = await client.post(f"{ADMIN}/templates/import-word", json=payload, headers=H(token))
    assert r.status_code == 200, r.text
    r = await client.post(f"{ADMIN}/templates/import-word", json=payload, headers=H(token))
    assert r.status_code == 409, r.text


async def test_import_create_requires_super_admin(client, clean_db):
    a = (await client.post(f"{ADMIN}/templates/import-word/analyze",
                           json={"file_name": "sample.docx", "content_base64": b64(make_sample_docx())},
                           headers=H((await create_admin())[1]))).json()
    _, atoken = await create_admin(role="admin")
    r = await client.post(f"{ADMIN}/templates/import-word",
                          json={"id": "x", "name_en": "X", "name_gu": "ક્સ", "content_gu": CONTENT_GU,
                                "fields": reviewed_fields(), "settings": a["settings"]},
                          headers=H(atoken))
    assert r.status_code == 403, r.text


async def test_generated_pdf_uses_configured_page_size(client, clean_db):
    admin, token = await create_admin()
    a = (await client.post(f"{ADMIN}/templates/import-word/analyze",
                           json={"file_name": "sample.docx", "content_base64": b64(make_sample_docx())},
                           headers=H(token))).json()
    await client.post(f"{ADMIN}/templates/import-word",
                      json={"id": "pdf_page_check", "name_en": "PDF Check", "name_gu": "પીડીએફ",
                            "content_gu": CONTENT_GU, "fields": reviewed_fields(),
                            "settings": settings_from_analysis(a)},
                      headers=H(token))
    await client.post(f"{ADMIN}/templates/pdf_page_check/publish", headers=H(token))

    _, utoken = await create_lawyer()
    values = {"applicant_name": "રમેશભાઈ પટેલ", "court": "અમદાવાદ", "case_number": "C/123/2025",
              "applicant_role": "અરજદાર", "opposite_role": "પ્રતિવાદી", "case_status": "ચાલુ"}
    r = await client.post(f"{API}/applications/preview",
                          json={"template_id": "pdf_page_check", "language": "gu", "values": values},
                          headers=H(utoken))
    assert r.status_code == 200, r.text
    r = await client.post(f"{API}/applications/download",
                          json={"template_id": "pdf_page_check", "language": "gu", "format": "pdf", "values": values},
                          headers=H(utoken))
    assert r.status_code == 200, r.text
    doc = r.json()
    assert doc["mime_type"] == "application/pdf"
    pdf_bytes = base64.b64decode(doc["base64"])
    assert pdf_bytes[:5] == b"%PDF-"
    # A4 MediaBox = 595 x 842 points (reportlab/playwright both write it plain)
    m = re.search(rb"/MediaBox\s*\[\s*([^\]]+)\]", pdf_bytes)
    assert m, f"MediaBox not found in PDF ({len(pdf_bytes)} bytes)"
    parts = [float(x) for x in m.group(1).split()]
    w, h = parts[2], parts[3]
    assert abs(w - 595) < 3 and abs(h - 842) < 3, f"expected A4, got {w}x{h}"


# ---------------------------------------------------------------------------
# Publish guard (Phase 0 root cause)
# ---------------------------------------------------------------------------

async def test_publish_malformed_record_returns_readable_error(client, clean_db):
    """A partial DB record (no name/content/fields) must 400, never 500."""
    admin, token = await create_admin()
    malformed = {
        "id": "malformed_partial_record",
        "status": "draft",
        "version": 1,
        "locked": False,
        "source": "admin_edited",
        "created_at": now().isoformat(),
        "updated_at": now().isoformat(),
    }
    await db.templates.insert_one(malformed)
    r = await client.post(f"{ADMIN}/templates/malformed_partial_record/publish", headers=H(token))
    assert r.status_code == 400, r.text
    detail = r.json()["detail"]
    assert "incomplete" in detail
    assert "name_en" in detail
    # Nothing was published / versioned
    assert await db.template_versions.count_documents({}) == 0
    doc = await db.templates.find_one({"id": "malformed_partial_record"})
    assert doc["status"] == "draft"


async def test_publish_malformed_fields_returns_readable_error(client, clean_db):
    admin, token = await create_admin()
    bad = {
        "id": "bad_fields_record",
        "name_en": "Bad",
        "name_gu": "બેડ",
        "content_en": "",
        "content_gu": "કંઈક",
        "status": "draft",
        "version": 1,
        "locked": False,
        "source": "admin_edited",
        "created_at": now().isoformat(),
        "updated_at": now().isoformat(),
    }
    await db.templates.insert_one(bad)  # no 'fields' key at all
    r = await client.post(f"{ADMIN}/templates/bad_fields_record/publish", headers=H(token))
    assert r.status_code == 400, r.text
    assert "field list" in r.json()["detail"]


async def test_valid_draft_still_publishes(client, clean_db):
    """The guard must not break publishing of valid templates."""
    admin, token = await create_admin()
    a = (await client.post(f"{ADMIN}/templates/import-word/analyze",
                           json={"file_name": "sample.docx", "content_base64": b64(make_sample_docx())},
                           headers=H(token))).json()
    await client.post(f"{ADMIN}/templates/import-word",
                      json={"id": "valid_publish_check", "name_en": "Valid", "name_gu": "વેલિડ",
                            "content_gu": CONTENT_GU, "fields": reviewed_fields(),
                            "settings": settings_from_analysis(a)},
                      headers=H(token))
    r = await client.post(f"{ADMIN}/templates/valid_publish_check/publish", headers=H(token))
    assert r.status_code == 200, r.text
    assert r.json()["template"]["status"] == "published"
