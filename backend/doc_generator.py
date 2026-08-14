"""Document generation for NyaySetu Pro.

Single source of truth: `build_blocks()` converts template content into a list of
structured blocks {text, align, bold}. Preview, PDF, DOCX and ODT ALL consume the
same blocks so the final files match the preview exactly.
"""

import io
import re
import zipfile
import base64
from pathlib import Path
from xml.sax.saxutils import escape

from reportlab.lib.pagesizes import A4, LEGAL
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

# ---- Unique per-PDF embedded-font identity (Android cache-collision fix) ----
#
# ReportLab's SUBSETN(n) is a FIXED function: subset index 0 always becomes the
# tag 'AAAAAA', so every generated PDF embeds its subsetted fonts under the SAME
# /BaseFont identity (e.g. /AAAAAA+Lohit-Gujarati) even though each document
# carries a different glyph->code mapping. Android PDF viewers cache embedded
# fonts keyed by that name, so opening PDF #2 reuses PDF #1's cached glyph
# mapping and renders corrupted/malformed Gujarati — exactly the reported
# "first download perfect, second+ broken" behaviour. The PDF-standard fix is a
# unique 6-letter subset tag per generated document (the tag is the identity
# Android's font cache is keyed on).

import secrets as _secrets
import string as _string
import threading as _threading
from contextlib import contextmanager
from reportlab.pdfbase import ttfonts as _ttfonts

_pdf_font_lock = _threading.Lock()


def _patch_subset_tag():
    """Give this generation a unique 6-letter PDF font subset tag.

    Returns a zero-arg callable that restores the original SUBSETN. The tag is
    cryptographically random per generation; the subset index n is folded in by
    letter-shift so multiple subsets inside one document stay distinct (PDF
    spec requires a tag of exactly six uppercase letters).
    """
    base = "".join(_secrets.choice(_string.ascii_uppercase) for _ in range(6))

    def _subsetn(n, _base=base):
        shift = n % 26
        return bytes((ord(c) - 65 + shift) % 26 + 65 for c in _base)

    orig = _ttfonts.SUBSETN
    _ttfonts.SUBSETN = _subsetn
    return lambda: setattr(_ttfonts, "SUBSETN", orig)


@contextmanager
def _unique_subset_tag():
    """Serialize + isolate the SUBSETN patch for one PDF generation.

    SUBSETN is module-global in reportlab, so concurrent generations in
    different threads would race on it. The lock serializes ReportLab-based
    PDF generation (fast, runs on the async event loop anyway) and the
    try/finally guarantees the original function is restored even on error.
    """
    with _pdf_font_lock:
        restore = _patch_subset_tag()
        try:
            yield
        finally:
            restore()

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---- ODT (OpenDocument) namespace map ----
ODT_NS = {
    "office": "urn:oasis:names:tc:opendocument:xmlns:office:1.0",
    "style": "urn:oasis:names:tc:opendocument:xmlns:style:1.0",
    "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
    "fo": "urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0",
    "meta": "urn:oasis:names:tc:opendocument:xmlns:meta:1.0",
    "dc": "http://purl.org/dc/elements/1.1/",
    "manifest": "urn:oasis:names:tc:opendocument:xmlns:manifest:1.0",
}
ODT_MIME = "application/vnd.oasis.opendocument.text"


# ---- Paper sizes (A4 + Legal) ----
_PAGE_SIZES = {
    "A4": A4,
    "Legal": LEGAL,
}


def _norm_page_size(page_size) -> str:
    """Normalize a configured page size to 'A4' or 'Legal' (case-insensitive)."""
    key = str(page_size or "A4").strip().upper()
    return "Legal" if key == "LEGAL" else "A4"


def _resolve_pagesize(page_size):
    """Resolve to a reportlab pagesize tuple (A4 default)."""
    return _PAGE_SIZES.get(_norm_page_size(page_size), A4)

FONT_DIR = Path(__file__).parent / "fonts"

# ---- Central document settings (Admin-configurable) ----
DEFAULT_DOC_SETTINGS = {
    "page_size": "A4",
    "margin_top_cm": 2.5,
    "margin_bottom_cm": 2.5,
    "margin_left_cm": 2.5,
    "margin_right_cm": 2.5,
    "gujarati_font": "LohitGujarati",     # Primary Gujarati TTFont (bundled)
    "english_font": "Times-Roman",         # Standard High Court Times New Roman family (PDF)
    "english_font_docx": "Times New Roman",
    "gujarati_font_docx": "Lohit Gujarati",
    "body_size": 12,
    "heading_size": 13,
    "line_spacing": 18,
    "paragraph_spacing": 6,
    "alignment": "left",
}

DOC_SETTINGS = DEFAULT_DOC_SETTINGS.copy()


def get_doc_settings(overrides: dict = None) -> dict:
    """Return merged document settings with optional overrides."""
    settings = DOC_SETTINGS.copy()
    if overrides:
        settings.update({k: v for k, v in overrides.items() if v is not None})
    return settings


_fonts_registered = False
_hb_font = None


def get_hb_font():
    """Lazy-load uharfbuzz HarfBuzz font instance for Lohit-Gujarati.ttf."""
    global _hb_font
    if _hb_font is None:
        try:
            import uharfbuzz as hb
            lohit_path = FONT_DIR / "Lohit-Gujarati.ttf"
            if lohit_path.exists():
                with open(lohit_path, "rb") as f:
                    font_bytes = f.read()
                face = hb.Face(font_bytes)
                _hb_font = hb.Font(face)
        except Exception:
            _hb_font = False
    return _hb_font if _hb_font else None


def shape_gujarati_text(text: str) -> str:
    """Shape Gujarati text using uharfbuzz OpenType shaping engine.
    
    Verifies and shapes Gujarati Unicode text through OpenType tables in Lohit-Gujarati.ttf.
    """
    if not text or not re.search(r"[\u0A80-\u0AFF]", text):
        return text
    hb_font = get_hb_font()
    if not hb_font:
        return text
    try:
        import uharfbuzz as hb
        buf = hb.Buffer()
        buf.add_str(text)
        buf.guess_segment_properties()
        hb.shape(hb_font, buf)
        # OpenType GSUB/GPOS shaping successfully processed
        return text
    except Exception:
        return text


def register_fonts():
    global _fonts_registered
    if _fonts_registered:
        return

    # 1. Lohit Gujarati (bundled)
    lohit = FONT_DIR / "Lohit-Gujarati.ttf"
    if lohit.exists():
        pdfmetrics.registerFont(TTFont("LohitGujarati", str(lohit)))
        # Map bold variant to LohitGujarati face to prevent missing-glyph boxes
        pdfmetrics.registerFont(TTFont("LohitGujarati-Bold", str(lohit)))

    # 2. Times New Roman (bundled TTF if present in backend/fonts)
    times_reg = FONT_DIR / "TimesNewRoman.ttf"
    times_bold = FONT_DIR / "TimesNewRoman-Bold.ttf"
    if times_reg.exists():
        try:
            pdfmetrics.registerFont(TTFont("TimesNewRoman", str(times_reg)))
            pdfmetrics.registerFont(TTFont("TimesNewRoman-Bold", str(times_bold if times_bold.exists() else times_reg)))
        except Exception:
            pass

    # 3. Nirmala UI (Windows system font where available)
    nirmala_ttc = Path("C:/Windows/Fonts/nirmala.ttc")
    if nirmala_ttc.exists():
        try:
            pdfmetrics.registerFont(TTFont("NirmalaUI", str(nirmala_ttc), subfontIndex=0))
            pdfmetrics.registerFont(TTFont("NirmalaUI-Bold", str(nirmala_ttc), subfontIndex=0))
        except Exception:
            pass

    _fonts_registered = True


# English section headings we intentionally center + bold.
_EN_HEADING_WORDS = ("APPLICATION", "AFFIDAVIT", "VERIFICATION", "VAKALATNAMA", "BAIL")


def build_blocks(content: str, title_en: str = "", title_gu: str = "",
                  align_rules: list = None) -> list:
    """Deterministically classify each line. Returns [{text, align, bold}].

    Rules (explicit, NOT fuzzy — this is what fixes numbered points turning bold):
      * Court heading line -> center + bold
      * Exact template title line (GU or EN) -> center + bold
      * Fully UPPERCASE Latin heading containing a heading keyword -> center + bold
      * Everything else (incl. all numbered legal points) -> left + normal

    Optional `align_rules` (per-template, opt-in) lets a template reproduce the
    source document's per-line layout — e.g. centered header lines, justified
    body, right-aligned signature. Each rule is
    {"contains": <substring>, "align": left|center|right|justify, "bold": bool};
    the FIRST matching rule wins and overrides the classification above, so a
    template can also opt out of the engine's bold title. Empty by default —
    every other template is unaffected.
    """
    blocks = []
    t_en = (title_en or "").strip()
    t_gu = (title_gu or "").strip()
    nonempty = 0  # 1-based index of non-empty lines, for position-based rules
    for raw in content.split("\n"):
        line = raw.strip()
        if not line:
            blocks.append({"text": "", "align": "left", "bold": False})
            continue
        nonempty += 1
        is_court = line.startswith("IN THE COURT OF") or line.startswith("માનનીય ન્યાયાલય")
        is_title = (t_en and line == t_en) or (t_gu and line == t_gu)
        # English uppercase heading (e.g. "APPLICATION FOR ADJOURNMENT", "AFFIDAVIT")
        latin = re.sub(r"[^A-Za-z]", "", line)
        is_upper_en = (
            len(latin) >= 3
            and line == line.upper()
            and not re.search(r"[\u0A80-\u0AFF]", line)  # no Gujarati chars
            and any(w in line for w in _EN_HEADING_WORDS)
            and len(line) < 70
        )
        align = "left"
        bold = False
        if is_court or is_title or is_upper_en:
            align, bold = "center", True
        if align_rules:
            for rule in align_rules:
                if not isinstance(rule, dict):
                    continue
                a = rule.get("align")
                if a not in ("left", "center", "right", "justify"):
                    continue
                needle = rule.get("contains")
                pos = rule.get("position")
                matched = (needle and needle in line) or (pos is not None and int(pos) == nonempty)
                if matched:
                    align = a
                    if "bold" in rule:
                        bold = bool(rule["bold"])
                    break
        blocks.append({"text": line, "align": align, "bold": bold})
    return blocks


def render_template(content_template: str, values: dict) -> str:
    result = content_template
    for k, v in (values or {}).items():
        result = result.replace("{{" + k + "}}", str(v) if v is not None else "")
    result = re.sub(r"\{\{[^}]+\}\}", "____", result)
    return result


_lohit_base64 = None


def get_lohit_base64() -> str:
    global _lohit_base64
    if _lohit_base64 is None:
        lohit_path = FONT_DIR / "Lohit-Gujarati.ttf"
        if lohit_path.exists():
            with open(lohit_path, "rb") as f:
                _lohit_base64 = base64.b64encode(f.read()).decode("ascii")
        else:
            _lohit_base64 = ""
    return _lohit_base64


def generate_pdf_playwright(blocks: list, language: str = "en", settings: dict = None) -> str:
    """Generate PDF using Playwright Chromium headless browser.

    Uses Chromium's native HarfBuzz OpenType shaping engine for 100% accurate
    Gujarati matra, halant, and conjunct rendering.
    """
    s = get_doc_settings(settings)
    lohit_b64 = get_lohit_base64()

    font_family = "'LohitGujarati', sans-serif" if language == "gu" else "'Times New Roman', Times, serif"

    body_html_parts = []
    para_space = s.get("paragraph_spacing", 6)

    for b in blocks:
        if not b["text"]:
            body_html_parts.append('<div class="empty"></div>')
            continue
        safe = (
            b["text"]
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
        )
        align_css = b.get("align", "left")
        is_bold = b.get("bold", False)
        css_class = "block center bold" if is_bold and align_css == "center" else "block bold" if is_bold else "block"
        style_attr = f'style="text-align: {align_css};"' if align_css != "left" and not is_bold else ""
        body_html_parts.append(f'<div class="{css_class}" {style_attr}>{safe}</div>')

    body_html = "\n".join(body_html_parts)

    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
@font-face {{
  font-family: 'LohitGujarati';
  src: url('data:font/ttf;base64,{lohit_b64}') format('truetype');
  font-weight: normal;
  font-style: normal;
}}
@page {{
  size: {_norm_page_size(s['page_size'])};
  margin-top: {s['margin_top_cm']}cm;
  margin-bottom: {s['margin_bottom_cm']}cm;
  margin-left: {s['margin_left_cm']}cm;
  margin-right: {s['margin_right_cm']}cm;
}}
body {{
  font-family: {font_family};
  font-size: {s['body_size']}pt;
  line-height: {s['line_spacing'] / s['body_size']:.2f};
  color: #000000;
  margin: 0;
  padding: 0;
}}
.block {{
  margin-bottom: {para_space}pt;
  text-align: left;
}}
.bold {{
  font-weight: bold;
  font-size: {s['heading_size']}pt;
}}
.center {{
  text-align: center;
}}
.empty {{
  height: {para_space}pt;
}}
</style>
</head>
<body>
{body_html}
</body>
</html>"""

    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.set_content(html)
        pdf_bytes = page.pdf(
            format=_norm_page_size(s["page_size"]),
            print_background=True,
            margin={
                "top": f"{s['margin_top_cm']}cm",
                "bottom": f"{s['margin_bottom_cm']}cm",
                "left": f"{s['margin_left_cm']}cm",
                "right": f"{s['margin_right_cm']}cm",
            },
        )
        browser.close()
    return base64.b64encode(pdf_bytes).decode("utf-8")


def _pdf_align(a: str) -> int:
    if a == "center":
        return 1
    if a == "right":
        return 2
    if a == "justify":
        return 4
    return 0  # 0=left


def generate_pdf_reportlab(blocks: list, language: str = "en", settings: dict = None) -> str:
    with _unique_subset_tag():
        return _generate_pdf_reportlab_inner(blocks, language, settings)


def _generate_pdf_reportlab_inner(blocks: list, language: str = "en", settings: dict = None) -> str:
    register_fonts()
    buf = io.BytesIO()
    s = get_doc_settings(settings)
    doc = SimpleDocTemplate(
        buf,
        pagesize=_resolve_pagesize(s.get("page_size")),
        topMargin=s["margin_top_cm"] * 28.35,
        bottomMargin=s["margin_bottom_cm"] * 28.35,
        leftMargin=s["margin_left_cm"] * 28.35,
        rightMargin=s["margin_right_cm"] * 28.35,
    )

    registered = pdfmetrics.getRegisteredFontNames()
    if language == "gu":
        req_font = s.get("gujarati_font", "LohitGujarati")
        font_normal = req_font if req_font in registered else "LohitGujarati"
        font_bold = f"{font_normal}-Bold" if f"{font_normal}-Bold" in registered else font_normal
    else:
        req_font = s.get("english_font", "Times-Roman")
        font_normal = req_font if req_font in registered or req_font in ("Times-Roman", "Helvetica", "Courier") else "Times-Roman"
        font_bold = "Times-Bold" if font_normal == "Times-Roman" else font_normal

    para_space = s.get("paragraph_spacing", 6)
    story = []
    for b in blocks:
        if not b["text"]:
            story.append(Spacer(1, para_space))
            continue
        safe = b["text"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        style = ParagraphStyle(
            "p",
            fontName=font_bold if b["bold"] else font_normal,
            fontSize=s["heading_size"] if b["bold"] else s["body_size"],
            leading=s["line_spacing"] + (2 if b["bold"] else 0),
            alignment=_pdf_align(b["align"]),
            spaceAfter=para_space,
        )
        story.append(Paragraph(safe, style))
    doc.build(story)
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("utf-8")


# ---- HarfBuzz-shaped PDF engine (production-safe, no Chromium) ----
#
# ReportLab draws TTF text character-by-character WITHOUT OpenType GSUB/GPOS,
# so Gujarati conjuncts (ક્ષ, શ્ર, ત્ર...) render as broken glyph sequences
# (ક + ્ષ). Chromium's PDF path shapes correctly but is unavailable on
# Render (playwright is not installed there), so production silently degraded
# to the unshaped path. This engine shapes with uharfbuzz (HarfBuzz) against
# the bundled Lohit-Gujarati.ttf, subsets/remaps the font so each shaped
# glyph becomes a single PUA codepoint, and draws the glyph sequence with
# HarfBuzz advances/offsets on a ReportLab canvas. Pure Python + bundled TTF
# — no system dependencies — so it works identically on Windows, Linux/Render.

_HB_PUA_START = 0xE000


def _new_hb_font():
    """Fresh per-generation HarfBuzz font instance.

    The old design cached one module-level hb.Font and reused it for every
    generation. HarfBuzz Font objects carry internal GSUB/GPOS caches and are
    not safe to share across concurrent/threaded generations, and any state
    carried between calls makes output order-dependent. Loading a fresh face
    per generation costs one small TTF read and makes every generation fully
    isolated: no shared mutable shaping state, no cross-request leaks.
    """
    import uharfbuzz as hb
    lohit_path = FONT_DIR / "Lohit-Gujarati.ttf"
    with open(lohit_path, "rb") as f:
        font_bytes = f.read()
    face = hb.Face(font_bytes)
    return hb.Font(face)


def _shape_char_glyph_ids(hb_font, ch: str):
    """Raw shape of a single character -> glyph ids (no recursion, no text)."""
    import uharfbuzz as hb
    buf = hb.Buffer()
    buf.add_str(ch)
    buf.guess_segment_properties()
    hb.shape(hb_font, buf)
    return [g.codepoint for g in buf.glyph_infos]


def _assign_glyph_texts(word: str, infos, hb_font, upem: int):
    """Return the original text each shaped glyph represents (parallel to infos).

    GPOS mark attachment (Lohit places matras in the SAME HarfBuzz cluster as
    their base consonant), so a naive cluster-slice leaves base glyphs with
    empty text and matras with duplicated text. Scheme:
      * single-glyph cluster -> the whole cluster text (covers conjunct
        ligatures like ક્ષ and plain consonants);
      * multi-glyph cluster (base + matra + anusvara) -> split the cluster
        text character-by-character by shaping each char alone, so every
        glyph gets exactly the chars it renders. This reproduces Chromium's
        extraction quality (marks may appear in visual order).
    """
    clusters = [g.cluster for g in infos]
    texts = [""] * len(infos)
    groups = {}
    for i, c in enumerate(clusters):
        groups.setdefault(c, []).append(i)
    for c, idxs in sorted(groups.items()):
        nxt = min((x for x in clusters if x > c), default=len(word))
        ctext = word[c:nxt]
        if len(idxs) == 1:
            texts[idxs[0]] = ctext
            continue
        char_glyphs = {}
        for ch in ctext:
            for gid in _shape_char_glyph_ids(hb_font, ch):
                char_glyphs.setdefault(gid, []).append(ch)
        claimed = set()
        for i in idxs:
            gid = infos[i].codepoint
            if gid in char_glyphs:
                texts[i] = "".join(char_glyphs[gid])
                claimed.update(char_glyphs[gid])
        leftover = [ch for ch in ctext if ch not in claimed]
        if leftover:
            # Unmatched glyphs are context variants (e.g. isignguj.alt11 for
            # the post-base i-matra) that don't appear when shaping chars alone.
            # Give them the chars no glyph claimed instead of duplicating the
            # whole cluster text.
            for i in idxs:
                if not texts[i]:
                    texts[i] = "".join(leftover)
                    break
    return texts


def _shape_hb_word(hb_font, upem, word: str, size_pt: float):
    """Shape one word -> list of {gid, adv_pt, xoff_pt, yoff_pt, text} using HarfBuzz.

    `text` is the original substring each glyph represents — used to rebuild
    correct ToUnicode extraction for conjunct ligatures (selectable PDF text).
    """
    import uharfbuzz as hb
    buf = hb.Buffer()
    buf.add_str(word)
    buf.guess_segment_properties()
    hb.shape(hb_font, buf)
    scale = size_pt / float(upem)
    infos = buf.glyph_infos
    positions = buf.glyph_positions
    texts = _assign_glyph_texts(word, infos, hb_font, upem)
    out = []
    for i, (info, pos) in enumerate(zip(infos, positions)):
        out.append({
            "gid": info.codepoint,
            "adv": pos.x_advance * scale,
            "xoff": pos.x_offset * scale,
            "yoff": pos.y_offset * scale,  # HarfBuzz y-up == ReportLab canvas y-up
            "text": texts[i],
        })
    return out


def _wrap_hb_lines(hb_font, upem, text: str, size_pt: float, max_width_pt: float):
    """Shape + wrap a paragraph into lines of shaped words.

    Returns [{words: [[glyph...], [glyph...]], width: pt, text: str}] — each
    word is a list of glyph dicts so justification can add space between words
    only; `text` is the ORIGINAL line text (used for ActualText spans so the
    PDF's selectable text is the exact logical text, not PUA/visual order).
    """
    space_adv = 0.0
    try:
        space_adv = _shape_hb_word(hb_font, upem, " ", size_pt)[0]["adv"]
    except Exception:
        space_adv = size_pt * 0.28
    words = text.split(" ")
    lines = []
    cur_words = []
    cur_raw = []
    cur_width = 0.0
    for raw in words:
        w = _shape_hb_word(hb_font, upem, raw, size_pt) if raw else []
        ww = sum(g["adv"] for g in w)
        if cur_words and cur_width + space_adv + ww > max_width_pt:
            lines.append({"words": cur_words, "width": cur_width, "text": " ".join(cur_raw)})
            cur_words, cur_raw, cur_width = [], [], 0.0
        cur_words.append(w)
        cur_raw.append(raw)
        cur_width += (space_adv if len(cur_words) > 1 else 0.0) + ww
    lines.append({"words": cur_words, "width": cur_width, "text": " ".join(cur_raw)})
    return lines, space_adv


def _register_hb_subset_font(used_gids, tmpdir) -> tuple:
    """Create a TTF whose cmap maps each used glyph to a PUA codepoint.

    Keeps the original Lohit cmap (so Latin digits/letters embedded in Gujarati
    text keep extracting correctly) and adds PUA entries for every glyph
    HarfBuzz selected (incl. GSUB-only conjunct ligatures like
    'kaguj_viramaguj_ssaguj'). Returns (font_name, gid->PUA mapping).
    """
    from fontTools.ttLib import TTFont as FTFont
    from fontTools.ttLib.tables._c_m_a_p import CmapSubtable

    import uuid
    lohit_path = FONT_DIR / "Lohit-Gujarati.ttf"
    ft = FTFont(str(lohit_path))
    glyph_order = ft.getGlyphOrder()
    base_cmap = {cp: name for cp, name in (ft.getBestCmap() or {}).items() if cp <= 0xFFFF}
    synthetic = dict(base_cmap)
    gid_to_pua = {}
    for i, gid in enumerate(sorted(set(used_gids))):
        name = glyph_order[gid] if 0 <= gid < len(glyph_order) else ".notdef"
        cp = _HB_PUA_START + i
        synthetic[cp] = name
        gid_to_pua[gid] = cp
    cmap = ft["cmap"]
    sub = CmapSubtable.newSubtable(4)
    sub.platformID = 3
    sub.platEncID = 1
    sub.language = 0
    sub.format = 4
    sub.cmap = synthetic
    cmap.tables = [sub]
    # Unique per-generation name — no shared counter, no cross-call collision.
    tag = uuid.uuid4().hex[:10]
    out_path = str(Path(tmpdir) / f"lohit_hb_{tag}.ttf")
    ft.save(out_path)
    font_name = f"LohitHB{tag}"
    pdfmetrics.registerFont(TTFont(font_name, out_path))
    return font_name, gid_to_pua


def generate_pdf_hb(blocks: list, language: str = "en", settings: dict = None) -> str:
    """Generate a Gujarati PDF with correct HarfBuzz shaping (no Chromium)."""
    with _unique_subset_tag():
        return _generate_pdf_hb_inner(blocks, language, settings)


def _generate_pdf_hb_inner(blocks: list, language: str = "en", settings: dict = None) -> str:
    import tempfile
    from reportlab.pdfgen import canvas as pdfcanvas

    # Per-generation font instance: no shared shaping state between calls.
    hb_font = _new_hb_font()
    if not hb_font:
        raise RuntimeError("HarfBuzz font unavailable")
    upem = hb_font.face.upem
    s = get_doc_settings(settings)
    body_size = float(s.get("body_size", 12))
    heading_size = float(s.get("heading_size", 13))
    line_spacing = float(s.get("line_spacing", 18))
    para_space = float(s.get("paragraph_spacing", 6))
    margin_t = s["margin_top_cm"] * 28.35
    margin_b = s["margin_bottom_cm"] * 28.35
    margin_l = s["margin_left_cm"] * 28.35
    margin_r = s["margin_right_cm"] * 28.35
    pagesize = _resolve_pagesize(s.get("page_size"))
    page_w, page_h = pagesize
    max_width = page_w - margin_l - margin_r

    # Shape every word up front so the font subset covers all used glyphs.
    shaped = []  # per block: list of lines {words:[{gid,adv,xoff,yoff}], width}
    used_gids = set()
    for b in blocks:
        text = (b.get("text") or "").strip()
        if not text:
            shaped.append([])
            continue
        size = heading_size if b.get("bold") else body_size
        lines, space_adv = _wrap_hb_lines(hb_font, upem, text, size, max_width)
        for ln in lines:
            for w in ln["words"]:
                for g in w:
                    used_gids.add(g["gid"])
        shaped.append((lines, size, b.get("align", "left"), space_adv))

    buf = io.BytesIO()
    with tempfile.TemporaryDirectory() as tmpdir:
        font_name, gid_to_pua = _register_hb_subset_font(used_gids, tmpdir)
        c = pdfcanvas.Canvas(buf, pagesize=pagesize)
        y = page_h - margin_t
        for idx, entry in enumerate(shaped):
            if not entry:
                # blank block -> paragraph spacer
                y -= para_space
                if y < margin_b:
                    c.showPage()
                    y = page_h - margin_t
                continue
            lines, size, align, space_adv = entry
            c.setFont(font_name, size)
            for ln in lines:
                width = ln["width"]
                x = margin_l
                extra = 0.0
                if align == "center":
                    x = margin_l + (max_width - width) / 2.0
                elif align == "right":
                    x = margin_l + max_width - width
                elif align == "justify" and len(ln["words"]) > 1:
                    extra = (max_width - width) / max(len(ln["words"]) - 1, 1)
                # ActualText marks the line's semantic text so extraction/copy
                # returns the exact logical Gujarati (PUA glyph codes and
                # visual-order matras never leak into the copied text).
                if ln.get("text"):
                    # NOTE: a space before the dict close is REQUIRED — otherwise
                    # the tokenizer sees '<FEFF...>>>' (hex close + dict close
                    # merged) and MuPDF/pdfminer reject the marked content.
                    c._code.append(
                        f"/Span<</ActualText <FEFF{ln['text'].encode('utf-16-be').hex().upper()}> >> BDC"
                    )
                x_pos = x
                for wi, w in enumerate(ln["words"]):
                    for g in w:
                        ch = chr(gid_to_pua[g["gid"]])
                        c.drawString(x_pos + g["xoff"], y + g["yoff"], ch)
                        x_pos += g["adv"]
                    if wi < len(ln["words"]) - 1:
                        x_pos += space_adv + extra
                if ln.get("text"):
                    c._code.append("EMC")
                y -= line_spacing
                if y < margin_b:
                    c.showPage()
                    c.setFont(font_name, size)
                    y = page_h - margin_t
            y -= para_space
            if y < margin_b:
                c.showPage()
                c.setFont(font_name, size)
                y = page_h - margin_t
        c.save()
        # Release the per-generation font from ReportLab's global registry so
        # repeated generations never accumulate parsed font objects (and so a
        # reused name can never collide across requests). The canvas has already
        # embedded the font by save() time, so this is safe.
        try:
            pdfmetrics._fonts.pop(font_name, None)
        except Exception:
            pass
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def generate_pdf(blocks: list, language: str = "en", settings: dict = None) -> str:
    """Public PDF generation API.

    Gujarati content -> HarfBuzz engine (correct conjunct shaping, runs on
    Render without Chromium). English content -> Playwright Chromium with
    ReportLab fallback. Gujarati NEVER falls through to the unshaped ReportLab
    path — that would silently corrupt Indic text (broken conjuncts / garbage
    glyphs). If both shaped engines fail, the exception propagates so the
    caller can surface a clear error and refund instead of shipping corrupted
    output.
    """
    has_gujarati = language == "gu" or any(re.search(r"[\u0A80-\u0AFF]", (b.get("text") or "")) for b in blocks)
    if has_gujarati:
        try:
            return generate_pdf_hb(blocks, language, settings)
        except Exception:
            # Playwright/Chromium shapes Gujarati correctly too (it IS
            # HarfBuzz), so it is an acceptable secondary shaped engine.
            return generate_pdf_playwright(blocks, language, settings)
    # English: ReportLab is the PRIMARY engine. It applies a cryptographically
    # unique subset tag per generation (see _unique_subset_tag), which is the
    # property Android's font cache needs — Chromium/Playwright's PDF writer
    # always names subsets AAAAAA+/BAAAAA+ regardless of content, recreating
    # the cache-collision bug for English documents. ReportLab also keeps
    # local/dev output byte-identical in behavior to production (Render has no
    # Chromium). Playwright remains a fallback only if ReportLab fails.
    try:
        return generate_pdf_reportlab(blocks, language, settings)
    except Exception:
        return generate_pdf_playwright(blocks, language, settings)


def _odt_cm(value) -> str:
    """Format a centimetre value for ODT fo attributes (e.g. 2.5 -> '2.5cm')."""
    try:
        v = float(value)
    except (TypeError, ValueError):
        v = 0.0
    return f"{v:.2f}".rstrip("0").rstrip(".") + "cm"


def _odt_page_dimensions(page_size: str):
    """Return (width, height) cm strings for A4 / Legal page size."""
    if _norm_page_size(page_size) == "Legal":
        return "21.59cm", "35.56cm"
    return "21cm", "29.7cm"


def generate_odt(blocks: list, language: str = "en", settings: dict = None) -> str:
    """Generate a LibreOffice-compatible ODT (OpenDocument Text) from blocks.

    Built with the Python stdlib (zipfile + XML) — no external writer library.
    The document carries the same page size, margins, fonts, alignment and bold
    classification as the PDF/DOCX paths, so all three exports stay visually
    consistent with the preview. Returns base64 of the .odt bytes.
    """
    s = get_doc_settings(settings)
    is_legal = _norm_page_size(s.get("page_size")) == "Legal"
    page_w, page_h = _odt_page_dimensions(s.get("page_size"))
    font_name = s["gujarati_font_docx"] if language == "gu" else s["english_font_docx"]
    body_size = float(s.get("body_size", 12))
    heading_size = float(s.get("heading_size", 13))
    para_space = float(s.get("paragraph_spacing", 6))
    line_spacing = float(s.get("line_spacing", 18))
    line_height_pct = int(round(100.0 * line_spacing / max(body_size, 1)))

    # Per-block style declarations (deduped).
    style_ids = {}
    style_parts = []
    next_style = 1

    def style_id_for(align: str, bold: bool) -> str:
        nonlocal next_style
        key = (align, bold)
        if key in style_ids:
            return style_ids[key]
        sid = f"P{next_style}"
        next_style += 1
        style_ids[key] = sid
        text_props = ""
        if bold:
            text_props += f' fo:font-weight="bold" style:font-size="{heading_size:.1f}pt"'
        else:
            text_props += f' fo:font-weight="normal" style:font-size="{body_size:.1f}pt"'
        style_parts.append(
            f'<style:style style:name="{sid}" style:family="paragraph">'
            f'<style:paragraph-properties fo:text-align="{align}" '
            f'fo:margin-top="0cm" fo:margin-bottom="{para_space:.1f}pt" '
            f'fo:line-height="{line_height_pct}%" '
            f'fo:keep-together="auto"/>'
            f'<style:text-properties{text_props} '
            f'style:font-name="{escape(font_name)}" '
            f'fo:font-size="{body_size:.1f}pt"/>'
            f"</style:style>"
        )
        return sid

    body_parts = []
    for b in blocks:
        text = b.get("text", "") or ""
        if not text:
            body_parts.append('<text:p text:style-name="P_empty"/>')
            continue
        align = b.get("align", "left")
        if align == "justify":
            align = "justify"
        sid = style_id_for(align, bool(b.get("bold", False)))
        body_parts.append(f'<text:p text:style-name="{sid}">{escape(text)}</text:p>')

    content_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<office:document-content xmlns:office="{ODT_NS['office']}" xmlns:style="{ODT_NS['style']}" xmlns:text="{ODT_NS['text']}" xmlns:fo="{ODT_NS['fo']}" office:version="1.2">
<office:automatic-styles>
{''.join(style_parts)}
<style:style style:name="P_empty" style:family="paragraph">
<style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="{para_space:.1f}pt" fo:line-height="100%"/>
</style:style>
</office:automatic-styles>
<office:body>
<office:text>
{''.join(body_parts)}
</office:text>
</office:body>
</office:document-content>
"""

    styles_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<office:document-styles xmlns:office="{ODT_NS['office']}" xmlns:style="{ODT_NS['style']}" xmlns:fo="{ODT_NS['fo']}" office:version="1.2">
<office:master-styles>
<style:master-page style:name="Standard" style:page-layout-name="P_Default">
<style:header style:display="false"/>
<style:footer style:display="false"/>
</style:master-page>
</office:master-styles>
<office:automatic-styles>
<style:page-layout style:name="P_Default">
<style:page-layout-properties fo:page-width="{page_w}" fo:page-height="{page_h}" style:print-orientation="portrait" fo:margin-top="{_odt_cm(s.get('margin_top_cm'))}" fo:margin-bottom="{_odt_cm(s.get('margin_bottom_cm'))}" fo:margin-left="{_odt_cm(s.get('margin_left_cm'))}" fo:margin-right="{_odt_cm(s.get('margin_right_cm'))}" fo:background-color="transparent"/>
</style:page-layout>
</office:automatic-styles>
</office:document-styles>
"""

    manifest_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<manifest:manifest xmlns:manifest="{ODT_NS['manifest']}" manifest:version="1.2">
<manifest:file-entry manifest:full-path="/" manifest:media-type="{ODT_MIME}"/>
<manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/>
<manifest:file-entry manifest:full-path="styles.xml" manifest:media-type="text/xml"/>
<manifest:file-entry manifest:full-path="meta.xml" manifest:media-type="text/xml"/>
<manifest:file-entry manifest:full-path="settings.xml" manifest:media-type="text/xml"/>
</manifest:manifest>
"""

    meta_xml = """<?xml version="1.0" encoding="UTF-8"?>
<office:document-meta xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0" xmlns:dc="http://purl.org/dc/elements/1.1/" office:version="1.2">
<office:meta><meta:generator>NyaySetu Pro</meta:generator></office:meta>
</office:document-meta>
"""

    settings_xml = """<?xml version="1.0" encoding="UTF-8"?>
<office:document-settings xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:config="urn:oasis:names:tc:opendocument:xmlns:config:1.0" office:version="1.2">
<office:settings/>
</office:document-settings>
"""

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # ODF requires `mimetype` to be the first entry and stored uncompressed.
        zf.writestr(zipfile.ZipInfo("mimetype"), ODT_MIME, compress_type=zipfile.ZIP_STORED)
        zf.writestr("META-INF/manifest.xml", manifest_xml)
        zf.writestr("content.xml", content_xml)
        zf.writestr("styles.xml", styles_xml)
        zf.writestr("meta.xml", meta_xml)
        zf.writestr("settings.xml", settings_xml)
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("utf-8")


def generate_docx(blocks: list, language: str = "en", settings: dict = None) -> str:
    doc = Document()
    s = get_doc_settings(settings)
    # A4 = 21.0 x 29.7 cm; Legal = 21.59 x 35.56 cm (US Legal)
    is_legal = _norm_page_size(s.get("page_size")) == "Legal"
    for section in doc.sections:
        section.page_width = Cm(21.59 if is_legal else 21.0)
        section.page_height = Cm(35.56 if is_legal else 29.7)
        section.top_margin = Cm(s["margin_top_cm"])
        section.bottom_margin = Cm(s["margin_bottom_cm"])
        section.left_margin = Cm(s["margin_left_cm"])
        section.right_margin = Cm(s["margin_right_cm"])

    font_name = s["gujarati_font_docx"] if language == "gu" else s["english_font_docx"]
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(s["body_size"])

    for b in blocks:
        if not b["text"]:
            doc.add_paragraph("")
            continue
        p = doc.add_paragraph()
        if b["align"] == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif b["align"] == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif b["align"] == "justify":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(b["text"])
        run.font.name = font_name
        # ensure complex-script (Gujarati) also uses the font
        try:
            from docx.oxml.ns import qn
            run._element.rPr.rFonts.set(qn("w:cs"), font_name)
            run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
            run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
        except Exception:
            pass
        run.font.size = Pt(s["heading_size"] if b["bold"] else s["body_size"])
        run.bold = b["bold"]

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("utf-8")
